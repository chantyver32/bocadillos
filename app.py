import streamlit as st
import barcode
from barcode.writer import ImageWriter
from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Generador de Códigos Masivo", page_icon="🏷️")

st.title("🏷️ Generador de Códigos de Barras en Lote")
st.write("Ingresa los datos línea por línea usando el formato: **Nombre del producto, Código**")

# Entrada principal: Área de texto múltiple
texto_multilinea = st.text_area(
    "Datos de los códigos:", 
    placeholder="Pastel de Chocolate Champlitte, 3000100090013\nGalletas de Mantequilla, 3000100090014\nPan de Muerto, 3000100090015",
    height=200
)

# Opciones de formato
st.markdown("### Opciones de Formato")
col1, col2 = st.columns(2)
with col1:
    quitar_ceros = st.checkbox("Eliminar ceros iniciales innecesarios", value=True)
with col2:
    validar_16 = st.checkbox("Validar longitud de 16 dígitos", value=False)

if st.button("Generar Documento Word", type="primary"):
    if texto_multilinea.strip():
        try:
            # Inicializar documento Word
            doc = Document()
            doc.add_heading('Etiquetas de Códigos de Barras', 0)
            
            # Crear tabla con 2 columnas para organizar los códigos
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            # Procesar las líneas de texto
            lineas = texto_multilinea.strip().split('\n')
            
            TIPO_CODIGO = barcode.get_barcode_class('code128')
            opciones_imagen = {
                'module_width': 0.3, 
                'module_height': 8.0, 
                'font_size': 10,
                'text_distance': 4.0,
            }
            
            row_cells = None
            
            for idx, linea in enumerate(lineas):
                # Validar que la línea tenga el formato correcto con la coma
                if ',' not in linea:
                    st.warning(f"La línea '{linea}' no tiene el formato correcto (Falta la coma). Se omitirá.")
                    continue
                    
                nombre_producto, codigo_texto = linea.split(',', 1)
                nombre_producto = nombre_producto.strip()
                codigo_texto = codigo_texto.strip()
                
                # Limpieza y validación
                if quitar_ceros:
                    codigo_texto = codigo_texto.lstrip('0')
                    if not codigo_texto:
                        codigo_texto = "0"
                
                if validar_16 and len(codigo_texto) != 16:
                    st.warning(f"⚠️ Atención: El código '{codigo_texto}' tiene {len(codigo_texto)} caracteres en lugar de 16.")
                
                # Generar imagen del código de barras en memoria
                buffer_img = BytesIO()
                codigo_generado = TIPO_CODIGO(codigo_texto, writer=ImageWriter())
                codigo_generado.write(buffer_img, options=opciones_imagen)
                buffer_img.seek(0) # Regresar el puntero al inicio del archivo en memoria
                
                # Acomodar en la tabla (alternar entre columna 0 y 1)
                if idx % 2 == 0:
                    row_cells = table.add_row().cells
                    col_idx = 0
                else:
                    col_idx = 1
                    
                # Insertar texto e imagen en la celda correspondiente
                cell = row_cells[col_idx]
                parrafo = cell.paragraphs[0]
                parrafo.add_run(f"{nombre_producto}\n").bold = True
                run_img = parrafo.add_run()
                run_img.add_picture(buffer_img, width=Inches(2.5)) # Ajustar el ancho de la imagen
            
            # Guardar el documento de Word en memoria
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.success("¡Documento Word generado con éxito!")
            
            # Botón para descargar el Word
            st.download_button(
                label="⬇️ Descargar Documento con Etiquetas",
                data=doc_buffer,
                file_name="etiquetas_codigos.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"Ocurrió un error al generar el documento: {e}")
    else:
        st.error("Por favor, ingresa al menos un producto para generar el documento.")
        
